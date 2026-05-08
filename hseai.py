import anthropic
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

client = anthropic.Anthropic()

SYSTEM_HSE = """
Sos un especialista en HSE con 15 años de experiencia en Oil & Gas 
en Vaca Muerta, Neuquén, Argentina. Conocés las normas IAPG y los 
procedimientos de reporte de la Secretaría de Energía de Neuquén.
Respondés ÚNICAMENTE en JSON sin texto adicional.
"""

SYSTEM_ENTREVISTA = """
Sos un especialista en HSE que realiza entrevistas post-incidente 
en operaciones de Oil & Gas en Vaca Muerta, Argentina.
Tu rol es identificar qué información crítica falta para completar 
un reporte HSE formal según normas IAPG.
Respondés ÚNICAMENTE en JSON sin texto adicional.
Hacés preguntas claras, específicas y en español rioplatense.
"""

def detectar_info_faltante(descripcion):
    """Detecta qué información crítica falta en la descripción."""
    
    prompt = f"""
<descripcion_evento>
{descripcion}
</descripcion_evento>


<instruccion>
Analizá la descripción del evento HSE e identificá qué información 
crítica falta para completar un reporte formal según normas IAPG.

CAMPOS OBLIGATORIOS — solo estos 5 son imprescindibles:
1. Fecha y hora aproximada del evento
2. Ubicación específica (pad, locación o área)
3. Nombre completo y legajo del trabajador involucrado
4. Si hubo lesiones y su tipo (aunque sea "sin lesiones")
5. Qué actividad se estaba realizando cuando ocurrió

Todo lo demás es deseable pero NO obligatorio para el reporte inicial.
El técnico HSE puede completar el resto después.
Marcá como obligatorio: false a todo lo que no esté en esa lista de 5.

Solo pedí información que realmente falte — no repitas lo que ya 
está en la descripción.
</instruccion>


<formato_salida>
{{
  "informacion_disponible": ["dato1", "dato2"],
  "informacion_faltante": [
    {{
      "campo": "nombre del campo faltante",
      "pregunta": "pregunta específica para obtener ese dato",
      "obligatorio": true/false
    }}
  ],
  "suficiente_para_reporte": true/false
}}
</formato_salida>
"""

    response = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=1024,
        temperature=0,
        system=SYSTEM_ENTREVISTA,
        messages=[{"role": "user", "content": prompt}]
    )

    texto = response.content[0].text
    texto = texto.replace("```json", "").replace("```", "").strip()
    match = re.search(r'\{.*\}', texto, re.DOTALL)
    if not match:
        return None
    try:
        return json.loads(match.group())
    except json.JSONDecodeError:
        return None

def generar_reporte_hse(descripcion_completa):
    """Genera el reporte HSE formal a partir de la descripción completa."""

    prompt = f"""
<descripcion_evento>
{descripcion_completa}
</descripcion_evento>

<instruccion>
Generá un reporte HSE formal completo con toda la información disponible.
Clasificá el incidente según normas IAPG.
Analizá la causa raíz usando el método de los 5 Por Qué.

IMPORTANTE: En los valores del JSON usá SOLO comillas simples dentro 
del texto, nunca comillas dobles. Evitá caracteres especiales.
</instruccion>

<formato_salida>
{{
  "numero_reporte": "HSE-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
  "fecha_hora": "{datetime.now().strftime('%d/%m/%Y %H:%M')}",
  "tipo_incidente": "",
  "severidad": "Bajo / Medio / Alto / Critico",
  "descripcion_formal": "",
  "ubicacion": "",
  "personas_involucradas": "",
  "lesiones": "",
  "danos_materiales": "",
  "impacto_ambiental": "",
  "causa_inmediata": "",
  "causa_raiz": "",
  "cinco_porques": ["", "", "", "", ""],
  "acciones_correctivas": [
    {{"accion": "", "responsable": "", "plazo": ""}}
  ],
  "acciones_preventivas": [
    {{"accion": "", "responsable": "", "plazo": ""}}
  ],
  "requiere_investigacion_formal": true,
  "notificar_a": [],
  "lecciones_aprendidas": ""
}}
</formato_salida>
"""

    response = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=2048,
        temperature=0,
        system=SYSTEM_HSE,
        messages=[{"role": "user", "content": prompt}]
    )

    texto = response.content[0].text
    texto = texto.replace("```json", "").replace("```", "").strip()

    match = re.search(r'\{.*\}', texto, re.DOTALL)
    if not match:
        print(f"\n⚠️  Debug — No se encontró JSON en la respuesta:")
        print(texto[:300])
        return None

    json_texto = match.group()

    try:
        return json.loads(json_texto)
    except json.JSONDecodeError:
        # Intentar limpiar caracteres problemáticos
        try:
            # Reemplazar comillas dobles dentro de valores
            import re as re2
            # Limpiar saltos de línea dentro de strings JSON
            json_limpio = re2.sub(
                r':\s*"([^"]*)"',
                lambda m: ': "' + m.group(1).replace('\n', ' ').replace('\r', '') + '"',
                json_texto
            )
            return json.loads(json_limpio)
        except json.JSONDecodeError as e:
            print(f"\n⚠️  Debug — Error de JSON: {e}")
            print(f"   Primeros 300 caracteres:")
            print(json_texto[:300])
            return None



def generar_word_hse(datos, ruta_salida):
    """Genera el reporte Word profesional."""

    doc = Document()

    titulo = doc.add_heading("REPORTE DE INCIDENTE HSE", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_num.add_run(
        f"N° {datos['numero_reporte']}   |   "
        f"Fecha: {datos['fecha_hora']}"
    )
    run.font.size = Pt(11)
    run.font.bold = True

    doc.add_paragraph()

    doc.add_heading("1. Datos del Incidente", 1)
    tabla = doc.add_table(rows=6, cols=2)
    tabla.style = "Table Grid"
    campos = [
        ["Tipo de incidente",     datos.get("tipo_incidente", "—")],
        ["Severidad",             datos.get("severidad", "—")],
        ["Ubicación",             datos.get("ubicacion", "—")],
        ["Personas involucradas", datos.get("personas_involucradas", "—")],
        ["Lesiones",              datos.get("lesiones", "—")],
        ["Daños materiales",      datos.get("danos_materiales", "—")],
    ]
    for i, (campo, valor) in enumerate(campos):
        tabla.rows[i].cells[0].text = campo
        tabla.rows[i].cells[1].text = valor
        tabla.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_heading("2. Descripción del Evento", 1)
    doc.add_paragraph(datos.get("descripcion_formal", "—"))

    doc.add_paragraph()
    doc.add_heading("3. Impacto Ambiental", 1)
    doc.add_paragraph(datos.get("impacto_ambiental", "—"))

    doc.add_paragraph()
    doc.add_heading("4. Análisis de Causas", 1)
    doc.add_paragraph(f"Causa inmediata: {datos.get('causa_inmediata', '—')}").runs[0].bold = True
    doc.add_paragraph()
    doc.add_paragraph(f"Causa raíz: {datos.get('causa_raiz', '—')}").runs[0].bold = True

    cinco_porques = datos.get("cinco_porques", [])
    if any(cinco_porques):
        doc.add_paragraph()
        doc.add_heading("Análisis 5 Por Qués", 2)
        for i, porque in enumerate(cinco_porques):
            if porque:
                doc.add_paragraph(f"¿Por qué {i+1}? {porque}", style="List Number")

    doc.add_paragraph()
    doc.add_heading("5. Acciones Correctivas", 1)
    tbl_ac = doc.add_table(rows=1, cols=3)
    tbl_ac.style = "Table Grid"
    for j, h in enumerate(["Acción", "Responsable", "Plazo"]):
        tbl_ac.rows[0].cells[j].text = h
        tbl_ac.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for ac in datos.get("acciones_correctivas", []):
        row = tbl_ac.add_row()
        row.cells[0].text = ac.get("accion", "—")
        row.cells[1].text = ac.get("responsable", "—")
        row.cells[2].text = ac.get("plazo", "—")

    doc.add_paragraph()
    doc.add_heading("6. Acciones Preventivas", 1)
    tbl_ap = doc.add_table(rows=1, cols=3)
    tbl_ap.style = "Table Grid"
    for j, h in enumerate(["Acción", "Responsable", "Plazo"]):
        tbl_ap.rows[0].cells[j].text = h
        tbl_ap.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for ap in datos.get("acciones_preventivas", []):
        row = tbl_ap.add_row()
        row.cells[0].text = ap.get("accion", "—")
        row.cells[1].text = ap.get("responsable", "—")
        row.cells[2].text = ap.get("plazo", "—")

    doc.add_paragraph()
    doc.add_heading("7. Lecciones Aprendidas", 1)
    doc.add_paragraph(datos.get("lecciones_aprendidas", "—"))

    notificar = datos.get("notificar_a", [])
    if notificar:
        doc.add_paragraph()
        doc.add_heading("8. Notificaciones Requeridas", 1)
        for n in notificar:
            doc.add_paragraph(n, style="List Bullet")

    doc.add_paragraph()
    requiere = datos.get("requiere_investigacion_formal", False)
    p_inv = doc.add_paragraph()
    run_inv = p_inv.add_run(
        f"¿Requiere investigación formal? {'SÍ' if requiere else 'NO'}"
    )
    run_inv.bold = True
    run_inv.font.size = Pt(12)
    if requiere:
        run_inv.font.color.rgb = RGBColor(220, 53, 69)

    doc.add_paragraph()
    pie = doc.add_paragraph(
        "Reporte generado por OilAI HSE Module — Vaca Muerta, Neuquén · 2026"
    )
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.runs[0].font.size = Pt(8)
    pie.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.save(ruta_salida)
    print(f"\n  ✅ Reporte Word guardado: {ruta_salida}")

