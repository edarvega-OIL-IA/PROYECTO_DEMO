import anthropic
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

client = anthropic.Anthropic()

# ── Categorías de causas raíz (formato Peduzzi/Shell) ──
CAUSAS_RAIZ = {
    "FACTORES PERSONALES": {
        "1. CAPACIDAD FÍSICA / MENTAL INADECUADA": [
            "1.1. Lesión y/o Enfermedad Preexistente",
            "1.2. Deficiencia Visual / Auditiva / Respiratoria",
            "1.3. Alteración provocada por Droga, Alcohol y/o Medicamentos",
            "1.4. Tamaño, dimensiones corporales inadecuadas",
            "1.5. Fatiga Física y/o Mental",
            "1.6. Otra discapacidad física/mental temporal o permanente"
        ],
        "2. COMPORTAMIENTO/MOTIVACIÓN INADECUADA": [
            "2.1. Presión percibida por el empleado",
            "2.2. Recompensa del comportamiento indebido",
            "2.3. Bajo nivel de Disciplina",
            "2.4. Actividades degradantes o poco motivadoras",
            "2.5. Ejemplo inapropiado del supervisor",
            "2.6. Otro comportamiento inapropiado"
        ],
        "3. NIVEL DE HABILIDAD DEFICIENTE": [
            "3.1. Evaluación inadecuada de las habilidades requeridas",
            "3.2. Entrenamiento práctico de las habilidades insuficiente",
            "3.3. Habilidades practicadas con poca frecuencia"
        ],
        "4. FALTA DE CONOCIMIENTO / CAPACITACIÓN": [
            "4.1. Falta de Capacitación",
            "4.2. Falta de conocimiento y/o experiencia",
            "4.3. Aplicación inadecuada de la capacitación recibida",
            "4.4. Evaluación deficiente de necesidades de capacitación",
            "4.5. Capacitación inadecuada y/o Deficiente"
        ]
    },
    "FACTORES DEL TRABAJO": {
        "5. FALLAS DE SUPERVISIÓN, LIDERAZGO Y/O PLANIFICACIÓN": [
            "5.1. Líneas de reporte, roles y responsabilidades poco claras",
            "5.2. Delegación de autoridad inadecuada o insuficiente",
            "5.3. Asignación de tareas inadecuada",
            "5.4. Supervisión directa insuficiente",
            "5.5. Habilidades de Liderazgo insuficientes",
            "5.6. Planificación inadecuada del trabajo",
            "5.7. Identificación inadecuada de Riesgos",
            "5.8. Gestión inadecuada de cambios",
            "5.9. Corrección inadecuada de actos y/o Condiciones inseguras"
        ],
        "6. INGENIERÍA, DISEÑO Y/O MANTENIMIENTO INADECUADOS": [
            "6.1. Diseño Técnico, Normas y/o Especificaciones inadecuados",
            "6.2. Evaluación inadecuada o insuficiente de Fallas",
            "6.3. Evaluación insuficiente de modificaciones",
            "6.4. Mantenimiento preventivo o correctivo Inadecuado",
            "6.5. Planificación inadecuada del uso de equipamiento"
        ],
        "7. ABASTECIMIENTO Y CONTROL DE PRODUCTOS Y SERVICIOS": [
            "7.1. Recepción y aceptación de materiales incorrectos",
            "7.2. Manipuleo y/o almacenamiento inadecuado de materiales",
            "7.3. Disponibilidad insuficiente de materiales o equipamiento",
            "7.4. Especificaciones deficientes en cuanto a los pedidos",
            "7.5. Evaluación de contratistas inadecuada",
            "7.6. Eliminación y Reemplazo inapropiado de piezas defectuosas"
        ],
        "8. PROCEDIMIENTOS / ESTÁNDARES DE TRABAJO INADECUADOS": [
            "8.1. Falta de procedimientos/Normas/estándares para la tarea",
            "8.2. Cumplimiento inadecuado o No cumplimiento de procedimientos",
            "8.3. Inadecuada difusión o implementación de procedimientos",
            "8.4. Uso indebido de Herramientas y/o equipo"
        ],
        "9. COMUNICACIÓN INADECUADA": [
            "9.1. Comunicación horizontal (entre pares) inadecuada",
            "9.2. Comunicación vertical (supervisor-subalterno) inadecuada",
            "9.3. Comunicación inadecuada entre grupos de trabajo"
        ],
        "10. REGLAS ORGANIZACIONALES": [
            "10.1. Metas y objetivos conflictivos o contradictorios",
            "10.2. Conducta inapropiada permitida o tolerada",
            "10.3. Políticas de sanciones y medidas disciplinarias poco claras"
        ]
    }
}

# ── Normativa nacional aplicable ──
NORMATIVA = """
NORMATIVA NACIONAL HSE ARGENTINA:
- Ley 19587/1972: Higiene y Seguridad en el Trabajo. Aplicable a todo establecimiento del territorio nacional. Dec. Reg. 351/79.
- Ley 24557/1995: Riesgos del Trabajo. Obligatoriedad de ART. Dec. Reg. 170/96.
- Decreto 1338/1996: Servicios de Medicina e Higiene y Seguridad en el Trabajo.
- Res. SRT 230/2003: Denuncia obligatoria de incidentes graves y mortales a la ART y SRT.
- Res. SRT 299/2011: Elementos de Protección Personal (EPP).
- Res. SRT 295/2003: Ergonomía y radiaciones.
- Res. SRT 592/2004: Trabajos con tensión eléctrica.
- Res. SRT 103/2005: Sistemas de Gestión de Seguridad y Salud.
- Res. SRT 801/2005: Sistema Globalmente Armonizado (SGA) de productos químicos.
- Res. SRT 37/2010 y 81/2019: Declaración anual de Agentes de Riesgo (RAR).
- Res. SRT 20/2018: Programa de Prevención para PyMEs.
- Decreto 911/1996: Construcción. Decreto 617/1997: Agraria. Decreto 249/2007: Minería.
- Res. SRT 84/2012: Protocolo iluminación. Res. SRT 85/2012: Protocolo ruido.
- Res. SRT 861/2015: Contaminantes químicos. Res. SRT 886/2015: Ergonomía.
- Res. SRT 900/2015: Puesta a tierra.
"""

SYSTEM_HSE_V2 = f"""
Sos un especialista en HSE (Salud, Seguridad y Medio Ambiente) con 
15 años de experiencia en Oil & Gas en Vaca Muerta, Neuquén, Argentina.
Conocés en profundidad los formatos de reporte de Shell, YPF, TotalEnergies 
y empresas contratistas como Peduzzi, OPS y San Antonio.

TERMINOLOGÍA CORRECTA Y VIGENTE:
- Todo evento que ocurre durante la tarea laboral se llama INCIDENTE.
- La palabra ACCIDENTE solo se usa para incidentes de tránsito no laborales.
- NUNCA usar la palabra accidente para eventos laborales.

CLASIFICACIÓN ACTUAL DE INCIDENTES:
1. Incidente Leve: sin lesiones o lesión menor tratada en campo
2. Incidente Grave: lesión que requiere atención médica fuera del lugar
3. Incidente con Pérdida de Personas: con días perdidos o fatalidad
4. Incidente con Pérdida Material: daño a equipos o instalaciones
5. Incidente Ambiental: liberación de fluidos o contaminación al ambiente
6. Incidente Vehicular: vehículos en operación (dentro del trabajo)
7. Near Miss / Casi-Incidente: evento sin consecuencias pero con potencial de daño
8. Condición Insegura: situación de riesgo identificada sin evento

DOS TIPOS DE INFORMES:
1. INFORME PRELIMINAR: 
   - Solo hechos concretos — qué pasó, cuándo, dónde, quién
   - SIN especulaciones, SIN causas, SIN análisis
   - Se envía INMEDIATAMENTE después del evento
   - Describir la historia de lo que pasó, no investigar el porqué

2. INFORME DE INVESTIGACIÓN FINAL:
   - Se genera DESPUÉS del proceso de investigación formal
   - Incluye análisis de causa raíz (espina de pescado, 5 porqués u otro método)
   - Las causas son REALES — no probables ni especulativas
   - Incluye plan de mitigación para evitar recurrencia
   - Lo produce la comisión de investigación

{NORMATIVA}

SISTEMA DE CAUSAS RAÍZ (10 categorías estándar Shell/Peduzzi):
Usás EXACTAMENTE las categorías del sistema de 10 grupos definido.

TUS REGLAS:
- Respondés ÚNICAMENTE en JSON sin texto adicional
- Nunca inventás datos — si falta información lo indicás como 'No especificado'
- NUNCA usás la palabra accidente para eventos laborales
- Citás la normativa aplicable según el tipo de incidente
- Usás SOLO comillas simples dentro de los valores de texto del JSON
"""

def detectar_info_faltante(descripcion):
    """Detecta qué información crítica falta."""
    prompt = f"""
<descripcion_evento>
{descripcion}
</descripcion_evento>

<instruccion>
Analizá la descripción del incidente HSE e identificá información faltante.

CAMPOS OBLIGATORIOS — solo estos 5:
1. Fecha y hora aproximada del evento
2. Ubicación específica (pad, locación o área)
3. Nombre completo y legajo del trabajador involucrado
4. Si hubo lesiones y su tipo (aunque sea 'sin lesiones')
5. Qué actividad se estaba realizando

Marcá como obligatorio: false todo lo que no esté en esa lista.
Solo pedí lo que realmente falte en la descripción.
</instruccion>

<formato_salida>
{{
  "informacion_disponible": [],
  "informacion_faltante": [
    {{
      "campo": "",
      "pregunta": "",
      "obligatorio": true
    }}
  ],
  "suficiente_para_reporte": true
}}
</formato_salida>
"""
    response = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=1024,
        temperature=0,
        system=SYSTEM_HSE_V2,
        messages=[{"role": "user", "content": prompt}]
    )
    texto = response.content[0].text
    texto = texto.replace("```json", "").replace("```", "").strip()
    match = re.search(r'\{.*\}', texto, re.DOTALL)
    if not match:
        return None
    try:
        return json.loads(match.group())
    except:
        return None

def generar_reporte_hse_v2(descripcion_completa):
    """Genera reporte HSE con formato Shell/Peduzzi actualizado."""

    causas_str = json.dumps(CAUSAS_RAIZ, ensure_ascii=False)

    prompt = f"""
    <descripcion_evento>
    {descripcion_completa}
    </descripcion_evento>

    <causas_raiz_disponibles>
        {causas_str}
    </causas_raiz_disponibles>

    <instruccion>
    Generá un reporte HSE completo usando terminología y clasificación vigente.

    REGLAS CRÍTICAS:
    1. NUNCA usar la palabra accidente para eventos laborales — siempre incidente
    2. Clasificar según el tipo de incidente vigente
    3. Elegir el formato según la gravedad:
        - preliminar: Near Miss, Condición Insegura, Incidente Leve
        - investigacion: Incidente Grave, con Pérdida de Personas, Ambiental, Vehicular

    PARA EL INFORME PRELIMINAR:
        - Describir SOLO los hechos
        - NO incluir causas ni análisis
        - causa_inmediata y causas_raiz_identificadas van como listas vacías []

    PARA EL INFORME DE INVESTIGACIÓN:
        - Incluir análisis completo de causa raíz
        - Las causas son REALES no probables
        - Incluir plan de mitigación concreto

    IMPORTANTE PARA EL JSON:
        - Usá SOLO comillas dobles en el JSON
        - NUNCA uses comillas simples dentro de valores de texto
        - NUNCA uses saltos de línea dentro de valores de string
        - Si necesitás una coma o punto dentro de un valor, escribila normalmente
        - El JSON debe ser perfectamente válido y parseable
    </instruccion>

    <formato_salida>
    Respondé ÚNICAMENTE con el JSON, sin texto antes ni después, sin bloques de código markdown.
    El JSON debe empezar con {{ y terminar con }}.

    {{
        "numero_reporte": "HSE-V2-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
        "fecha_hora": "{datetime.now().strftime('%d/%m/%Y %H:%M')}",
        "formato": "preliminar o investigacion",
        "tipo_incidente": "leve o grave o con_perdida_personas o con_perdida_material o ambiental o vehicular o near_miss o condicion_insegura",
        "clasificacion_gravedad": "Nivel 1 Leve o Nivel 2 Grave o Nivel 3 Con perdida",
        "circunstancia": "In Itinere o En Trayecto o En Tareas o Otros",
        "datos_involucrado": {{
            "apellido": "",
            "nombre": "",
            "dni": "",
            "edad": "",
            "funcion": "",
            "legajo": "",
            "antiguedad_empresa": "",
            "antiguedad_funcion": "",
            "personal_propio": true,
            "empresa_contratista": ""
    }},
    "servicio_obra": "",
    "cliente": "",
    "ubicacion": "",
    "descripcion_suceso": "",
    "descripcion_consecuencia": "",
    "acciones_inmediatas": "",
    "comunicaciones": {{
        "cliente": false,
        "art": false,
        "ambulancia": false,
        "bomberos": false,
        "policia": false
     }},
    "diagnostico_medico": "",
    "normativa_aplicable": [],
    "causa_inmediata": "",
    "causas_raiz_identificadas": [],
    "metodo_analisis_causa_raiz": "",
    "cinco_porques": [],
    "acciones_correctivas": [],
    "plan_mitigacion": [],
    "requiere_investigacion_formal": false,
    "requiere_denuncia_art": false,
    "entrevistas": [],
    "redactado_por": "OilAI HSE Pro v2",
    "lecciones_aprendidas": ""
    }}
    </formato_salida>
    """

    for intento in range(3):
        try:
            response = client.messages.create(
                model="claude-sonnet-4-5",
                max_tokens=3000,
                temperature=0,
                system=SYSTEM_HSE_V2,
                messages=[{"role": "user", "content": prompt}]
            )

            texto = response.content[0].text.strip()

            # Limpiar bloques markdown si los hay
            texto = texto.replace("```json", "").replace("```", "").strip()

            # Buscar el JSON
            match = re.search(r'\{.*\}', texto, re.DOTALL)
            if not match:
                continue

            json_texto = match.group()

            # Intentar parsear directo
            try:
                return json.loads(json_texto)
            except json.JSONDecodeError:
                pass

            # Limpiar saltos de línea dentro de strings
            def limpiar_string_json(m):
                contenido = m.group(1)
                contenido = contenido.replace('\n', ' ').replace('\r', '').replace('\t', ' ')
                return f'"{contenido}"'

            json_limpio = re.sub(r'"((?:[^"\\]|\\.)*)"', limpiar_string_json, json_texto)

            try:
                return json.loads(json_limpio)
            except json.JSONDecodeError:
                continue

        except Exception:
            continue

    return None




def generar_word_preliminar(datos, ruta_salida):
    """Genera informe preliminar estilo Shell RG-11-01."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Título
    titulo = doc.add_heading("INFORME PRELIMINAR DE INCIDENTE", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        f"N° {datos.get('numero_reporte', '—')}   |   "
        f"Fecha: {datos.get('fecha_hora', '—')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(10)

    # Aviso preliminar
    aviso = doc.add_paragraph()
    aviso_run = aviso.add_run(
        "⚠️ INFORME PRELIMINAR: Contiene únicamente hechos concretos. "
        "No incluye análisis de causas ni especulaciones. "
        "El Informe de Investigación Final se emitirá una vez completado el proceso de investigación."
    )
    aviso_run.font.size = Pt(9)
    aviso_run.font.color.rgb = RGBColor(180, 80, 0)
    aviso_run.bold = True

    doc.add_paragraph()

    # Datos principales
    doc.add_heading("1. DATOS DEL INCIDENTE", 1)
    tbl = doc.add_table(rows=4, cols=4)
    tbl.style = "Table Grid"
    datos_tbl = [
        ["N° Evento", datos.get("numero_reporte", "—"),
         "Fecha y Hora", datos.get("fecha_hora", "—")],
        ["Cliente", datos.get("cliente", "—"),
         "Servicio/Obra", datos.get("servicio_obra", "—")],
        ["Ubicación", datos.get("ubicacion", "—"),
         "Circunstancia", datos.get("circunstancia", "—")],
        ["Tipo de Incidente", datos.get("tipo_incidente", "—").replace("_", " ").upper(),
         "Gravedad", datos.get("clasificacion_gravedad", "—")],
    ]
    for i, fila in enumerate(datos_tbl):
        for j, val in enumerate(fila):
            cell = tbl.rows[i].cells[j]
            run = cell.paragraphs[0].add_run(val)
            if j % 2 == 0:
                run.bold = True
            run.font.size = Pt(9)

    doc.add_paragraph()

    # Datos del involucrado
    doc.add_heading("2. DATOS DEL INVOLUCRADO", 1)
    inv = datos.get("datos_involucrado", {})
    tbl2 = doc.add_table(rows=4, cols=4)
    tbl2.style = "Table Grid"
    datos_inv = [
        ["Apellido", inv.get("apellido", "—"), "D.N.I.", inv.get("dni", "—")],
        ["Nombre", inv.get("nombre", "—"), "Edad", inv.get("edad", "—")],
        ["Función", inv.get("funcion", "—"), "Legajo", inv.get("legajo", "—")],
        ["Antigüedad Empresa", inv.get("antiguedad_empresa", "—"),
         "Antigüedad Función", inv.get("antiguedad_funcion", "—")],
    ]
    for i, fila in enumerate(datos_inv):
        for j, val in enumerate(fila):
            cell = tbl2.rows[i].cells[j]
            run = cell.paragraphs[0].add_run(val)
            if j % 2 == 0:
                run.bold = True
            run.font.size = Pt(9)

    doc.add_paragraph()

    # Descripción del suceso — SOLO HECHOS
    doc.add_heading("3. DESCRIPCIÓN DEL SUCESO (Solo hechos)", 1)
    p_aviso2 = doc.add_paragraph()
    p_aviso2.add_run(
        "Descripción objetiva de lo ocurrido. Sin causas ni especulaciones."
    ).font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph(datos.get("descripcion_suceso", "—"))

    doc.add_paragraph()

    # Consecuencias
    doc.add_heading("4. DESCRIPCIÓN DE LA CONSECUENCIA", 1)
    doc.add_paragraph(datos.get("descripcion_consecuencia", "—"))

    doc.add_paragraph()

    # Diagnóstico médico
    if datos.get("diagnostico_medico"):
        doc.add_heading("5. DIAGNÓSTICO MÉDICO", 1)
        doc.add_paragraph(datos.get("diagnostico_medico", "No aplica"))
        doc.add_paragraph()

    # Acciones inmediatas
    doc.add_heading("6. ACCIONES INMEDIATAS TOMADAS", 1)
    doc.add_paragraph(datos.get("acciones_inmediatas", "—"))

    doc.add_paragraph()

    # Comunicaciones
    doc.add_heading("7. COMUNICACIONES REALIZADAS", 1)
    com = datos.get("comunicaciones", {})
    tbl3 = doc.add_table(rows=3, cols=4)
    tbl3.style = "Table Grid"
    tbl3.rows[0].cells[0].text = "Se comunicó al Cliente"
    tbl3.rows[0].cells[1].text = "SÍ" if com.get("cliente") else "NO"
    tbl3.rows[0].cells[2].text = "Intervino Ambulancia"
    tbl3.rows[0].cells[3].text = "SÍ" if com.get("ambulancia") else "NO"
    tbl3.rows[1].cells[0].text = "Se comunicó a la A.R.T."
    tbl3.rows[1].cells[1].text = "SÍ" if com.get("art") else "NO"
    tbl3.rows[1].cells[2].text = "Intervino Policía"
    tbl3.rows[1].cells[3].text = "SÍ" if com.get("policia") else "NO"
    tbl3.rows[2].cells[0].text = "Requiere denuncia ART"
    tbl3.rows[2].cells[1].text = "SÍ" if datos.get("requiere_denuncia_art") else "NO"
    tbl3.rows[2].cells[2].text = "Requiere investigación formal"
    tbl3.rows[2].cells[3].text = "SÍ" if datos.get("requiere_investigacion_formal") else "NO"

    doc.add_paragraph()

    # Normativa aplicable
    normativa = datos.get("normativa_aplicable", [])
    if normativa:
        doc.add_heading("8. NORMATIVA APLICABLE", 1)
        for n in normativa:
            doc.add_paragraph(n, style="List Bullet").runs[0].font.size = Pt(9)
        doc.add_paragraph()

    # Pie
    doc.add_paragraph()
    pie = doc.add_paragraph(
        f"Redactado por: {datos.get('redactado_por', 'OilAI HSE Pro v2')}   |   "
        "Vaca Muerta, Neuquén · 2026"
    )
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.runs[0].font.size = Pt(8)
    pie.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.save(ruta_salida)

def generar_word_investigacion(datos, ruta_salida):
    """Genera informe de investigación completo estilo Peduzzi F01 PG-14."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Título
    titulo = doc.add_heading("REPORTE DE INVESTIGACIÓN DE INCIDENTES", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        f"N° {datos.get('numero_reporte', '—')}   |   {datos.get('fecha_hora', '—')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True

    doc.add_paragraph()

    # SECCIÓN A — Información general
    doc.add_heading("SECCIÓN A: INFORMACIÓN GENERAL", 1)
    tbl_a = doc.add_table(rows=4, cols=4)
    tbl_a.style = "Table Grid"
    inv = datos.get("datos_involucrado", {})
    datos_a = [
        ["Obra/Servicio", datos.get("servicio_obra", "—"),
         "Lugar", datos.get("ubicacion", "—")],
        ["Fecha", datos.get("fecha_hora", "—").split()[0] if " " in datos.get("fecha_hora", "") else "—",
         "Hora", datos.get("fecha_hora", "—").split()[-1] if " " in datos.get("fecha_hora", "") else "—"],
        ["Cliente", datos.get("cliente", "—"),
         "Tipo de Incidente", datos.get("tipo_incidente", "—").replace("_", " ").upper()],
        ["Gravedad", datos.get("clasificacion_gravedad", "—"),
         "Circunstancia", datos.get("circunstancia", "—")],
    ]
    for i, fila in enumerate(datos_a):
        for j, val in enumerate(fila):
            cell = tbl_a.rows[i].cells[j]
            run = cell.paragraphs[0].add_run(val)
            if j % 2 == 0:
                run.bold = True
            run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph(
        f"Descripción de la situación: {datos.get('descripcion_suceso', '—')}"
    )

    # Datos del involucrado
    doc.add_paragraph()
    doc.add_heading("Datos del Involucrado", 2)
    tbl_inv = doc.add_table(rows=5, cols=2)
    tbl_inv.style = "Table Grid"
    campos_inv = [
        ["Nombre y Apellido", f"{inv.get('nombre', '—')} {inv.get('apellido', '—')}"],
        ["Función", inv.get("funcion", "—")],
        ["Legajo", inv.get("legajo", "—")],
        ["Antigüedad", inv.get("antiguedad_empresa", "—")],
        ["Diagnóstico Médico", datos.get("diagnostico_medico", "No aplica")],
    ]
    for i, (campo, valor) in enumerate(campos_inv):
        tbl_inv.rows[i].cells[0].paragraphs[0].add_run(campo).bold = True
        tbl_inv.rows[i].cells[1].text = valor

    doc.add_paragraph()

    # SECCIÓN B — Hallazgos
    doc.add_heading("SECCIÓN B: HALLAZGOS DEL INCIDENTE", 1)
    doc.add_heading("B.1. Descripción y Consecuencias", 2)
    doc.add_paragraph(datos.get("descripcion_consecuencia", "—"))

    # Entrevistas
    entrevistas = datos.get("entrevistas", [])
    if entrevistas and entrevistas[0].get("nombre"):
        doc.add_paragraph()
        doc.add_heading("B.2. Entrevistas al Personal Involucrado", 2)
        for i, ent in enumerate(entrevistas):
            if ent.get("nombre"):
                p = doc.add_paragraph()
                p.add_run(
                    f"Entrevista N°{i+1}: {ent.get('nombre', '—')} — {ent.get('cargo', '—')}"
                ).bold = True
                doc.add_paragraph(f"Empresa: {ent.get('empresa', '—')}   Fecha: {ent.get('fecha', '—')}")
                doc.add_paragraph(f"Declaración: \"{ent.get('declaracion', '—')}\"")
                doc.add_paragraph()

    # SECCIÓN C — Comité
    doc.add_heading("SECCIÓN C: COMITÉ DE INVESTIGACIÓN", 1)
    tbl_c = doc.add_table(rows=2, cols=2)
    tbl_c.style = "Table Grid"
    tbl_c.rows[0].cells[0].paragraphs[0].add_run("Responsable").bold = True
    tbl_c.rows[0].cells[1].text = datos.get("redactado_por", "OilAI HSE Pro v2")
    tbl_c.rows[1].cells[0].paragraphs[0].add_run("Fecha de cierre").bold = True
    tbl_c.rows[1].cells[1].text = datos.get("fecha_hora", "—").split()[0] if " " in datos.get("fecha_hora", "") else "—"

    doc.add_paragraph()

    # Listado global de causas raíz
    doc.add_heading("LISTADO GLOBAL DE CAUSAS RAÍZ", 2)
    causas_identificadas = [
        cr.get("codigo", "") if isinstance(cr, dict) else ""
        for cr in datos.get("causas_raiz_identificadas", [])]

    for grupo, subgrupos in CAUSAS_RAIZ.items():
        doc.add_paragraph(grupo).runs[0].bold = True
        for subgrupo, causas in subgrupos.items():
            p_sub = doc.add_paragraph(subgrupo, style="List Bullet")
            p_sub.runs[0].bold = True
            for causa in causas:
                partes = causa.split(".")
                codigo = f"{partes[0]}.{partes[1].split(' ')[0]}" if len(partes) > 1 else ""
                marcado = "☒" if any(codigo in c for c in causas_identificadas) else "☐"
                p = doc.add_paragraph(style="List Bullet 2")
                p.add_run(f"{marcado} {causa}").font.size = Pt(9)

    doc.add_paragraph()

    # SECCIÓN D — Conclusiones
    doc.add_heading("SECCIÓN D: CONCLUSIONES, CAUSAS Y ACCIONES", 1)
    doc.add_heading("Sumario del Evento (post-investigación)", 2)
    doc.add_paragraph(datos.get("descripcion_suceso", "—"))

    doc.add_paragraph()
    doc.add_heading("Causa Inmediata", 2)
    doc.add_paragraph(datos.get("causa_inmediata", "—"))

    doc.add_paragraph()
    doc.add_heading("Causas Raíz Identificadas (causas reales)", 2)
    p_nota = doc.add_paragraph()
    p_nota.add_run(
        "Nota: Las causas raíz son REALES — resultado del proceso de investigación. No son probables ni especulativas."
    ).font.color.rgb = RGBColor(128, 128, 128)
    for i, cr in enumerate(datos.get("causas_raiz_identificadas", []), 1):
        doc.add_paragraph(
            f"{i:02d}. [{cr.get('codigo', '—')}] {cr.get('descripcion', '—')}"
        )

    # Método análisis causa raíz
    metodo = datos.get("metodo_analisis_causa_raiz", "")
    if metodo:
        doc.add_paragraph()
        doc.add_heading(f"Método de Análisis: {metodo}", 2)
        cinco = datos.get("cinco_porques", [])
        if any(cinco):
            for i, porque in enumerate(cinco, 1):
                if porque:
                    doc.add_paragraph(f"¿Por qué {i}? {porque}", style="List Number")

    doc.add_paragraph()
    doc.add_heading("Acciones Correctivas / Preventivas", 2)
    tbl_d = doc.add_table(rows=1, cols=4)
    tbl_d.style = "Table Grid"
    for j, h in enumerate(["N°", "Descripción", "Responsable", "Plazo"]):
        tbl_d.rows[0].cells[j].paragraphs[0].add_run(h).bold = True

    for ac in datos.get("acciones_correctivas", []):
        if isinstance(ac, dict):
            row = tbl_d.add_row()
            row.cells[0].text = str(ac.get("numero", "—"))
            row.cells[1].text = ac.get("descripcion", "—")
            row.cells[2].text = ac.get("responsable", "—")
            row.cells[3].text = ac.get("plazo", "—")



    # Plan de mitigación
    plan = datos.get("plan_mitigacion", [])
    if plan and plan[0].get("accion"):
        doc.add_paragraph()
        doc.add_heading("Plan de Mitigación (para evitar recurrencia)", 2)
        tbl_pm = doc.add_table(rows=1, cols=5)
        tbl_pm.style = "Table Grid"
        for j, h in enumerate(["Acción", "Objetivo", "Responsable", "Plazo", "Indicador"]):
            tbl_pm.rows[0].cells[j].paragraphs[0].add_run(h).bold = True

        for pm in plan:
            if isinstance(pm, dict):
                row = tbl_pm.add_row()
                row.cells[0].text = pm.get("accion", "—")
                row.cells[1].text = pm.get("objetivo", "—")
                row.cells[2].text = pm.get("responsable", "—")
                row.cells[3].text = pm.get("plazo", "—")
                row.cells[4].text = pm.get("indicador", "—")


    doc.add_paragraph()

    # SECCIÓN E — Revisión gerencia
    doc.add_heading("SECCIÓN E: REVISIÓN DE GERENCIA / CONTROL", 1)
    preguntas_e = [
        ("¿Se completó la investigación en forma satisfactoria?", "Sí"),
        ("¿Fueron identificadas todas las causas raíces reales?", "Sí"),
        ("¿Se consideraron medidas correctivas para todas las causas?", "Sí"),
        ("¿Se elaboró plan de mitigación para evitar recurrencia?", "Sí"),
        ("¿Considera necesario profundizar en la investigación?",
         "Sí" if datos.get("requiere_investigacion_formal") else "No"),
        ("¿Se realizó denuncia a la ART?",
         "Sí" if datos.get("requiere_denuncia_art") else "No aplica"),
    ]
    tbl_e = doc.add_table(rows=len(preguntas_e), cols=2)
    tbl_e.style = "Table Grid"
    for i, (pregunta, resp) in enumerate(preguntas_e):
        tbl_e.rows[i].cells[0].text = pregunta
        tbl_e.rows[i].cells[1].text = resp

    # Normativa aplicable
    normativa = datos.get("normativa_aplicable", [])
    if normativa:
        doc.add_paragraph()
        doc.add_heading("Normativa Aplicable", 2)
        for n in normativa:
            doc.add_paragraph(n, style="List Bullet").runs[0].font.size = Pt(9)

    # Lecciones aprendidas
    doc.add_paragraph()
    doc.add_heading("Lecciones Aprendidas", 2)
    doc.add_paragraph(datos.get("lecciones_aprendidas", "—"))

    # Pie
    doc.add_paragraph()
    pie = doc.add_paragraph(
        "Reporte generado por OilAI HSE Pro v2 — "
        "Formato Shell/Peduzzi actualizado · Vaca Muerta, Neuquén · 2026"
    )
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.runs[0].font.size = Pt(8)
    pie.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.save(ruta_salida)

def generar_word_hse_v2(datos, ruta_salida):
    """Selecciona el formato según gravedad y genera el Word."""
    formato = datos.get("formato", "preliminar")
    if formato == "investigacion":
        generar_word_investigacion(datos, ruta_salida)
    else:
        generar_word_preliminar(datos, ruta_salida)