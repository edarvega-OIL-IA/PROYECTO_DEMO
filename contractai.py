import anthropic
import fitz
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from prompts import CONTRACTAI_SYSTEM, prompt_analisis_contrato
from utils import limpiar_json, verificar_respuesta, timestamp, separador

client = anthropic.Anthropic()

# ── Leer PDF ──
def leer_pdf(ruta):
    doc = fitz.open(ruta)
    paginas = []
    for i, pagina in enumerate(doc):
        texto = pagina.get_text()
        texto = texto.replace('\\', ' ').replace('\x00', '')
        texto = texto.replace('"', "'")  # ← esta línea faltaba
        paginas.append({"numero": i + 1, "texto": texto})
    doc.close()
    return paginas



# ── Dividir en chunks ──
def dividir_en_chunks(paginas, caracteres_por_chunk=50000):
    chunks = []
    chunk_actual = ""
    paginas_chunk = []
    for pagina in paginas:
        if len(chunk_actual) + len(pagina["texto"]) > caracteres_por_chunk:
            if chunk_actual:
                chunks.append({"texto": chunk_actual, "paginas": paginas_chunk.copy()})
            chunk_actual = pagina["texto"]
            paginas_chunk = [pagina["numero"]]
        else:
            chunk_actual += pagina["texto"]
            paginas_chunk.append(pagina["numero"])
    if chunk_actual:
        chunks.append({"texto": chunk_actual, "paginas": paginas_chunk.copy()})
    return chunks

# ── Analizar chunk en texto plano ──
def analizar_chunk(chunk, numero, total):
    print(f"  Analizando parte {numero}/{total} "
          f"(págs. {chunk['paginas'][0]}-{chunk['paginas'][-1]})...")
    prompt = f"""
<parte_documento>
Parte {numero} de {total} — Páginas {chunk['paginas'][0]} a {chunk['paginas'][-1]}
{chunk['texto']}
</parte_documento>

<instruccion>
Analizá esta sección. Respondé en texto plano con este formato exacto:

OBLIGACIONES:
- [obligación]

RIESGOS:
- [riesgo]

CLAUSULAS PROBLEMATICAS:
- [cláusula]

Si no hay contenido relevante escribí: ninguna identificada.
</instruccion>
"""
    try:
        response = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=2048,
            temperature=0,
            system=CONTRACTAI_SYSTEM,
            messages=[{"role": "user", "content": prompt}]
        )
        return {
            "parte": numero,
            "paginas": f"{chunk['paginas'][0]}-{chunk['paginas'][-1]}",
            "analisis": response.content[0].text,
            "exitoso": True
        }
    except Exception as e:
        return {
            "parte": numero,
            "paginas": f"{chunk['paginas'][0]}-{chunk['paginas'][-1]}",
            "error": str(e),
            "exitoso": False
        }

# ── Consolidar análisis de chunks ──
def consolidar(resultados):
    print("  Consolidando análisis...")
    resumen = ""
    for r in resultados:
        resumen += f"\n=== PARTE {r['parte']} (págs. {r['paginas']}) ===\n"
        resumen += r["analisis"] + "\n"

    prompt = f"""
<analisis_por_partes>
{resumen}
</analisis_por_partes>

<instruccion>
Consolidá los análisis parciales en un resumen ejecutivo final.
Eliminá duplicados y priorizá los puntos más importantes.
IMPORTANTE: en los valores del JSON usá SOLO comillas simples 
dentro del texto, nunca comillas dobles.
Respondé ÚNICAMENTE con este JSON sin texto adicional.
</instruccion>

<formato_salida>
{{
  "tipo_documento": "tipo aqui",
  "partes": {{"contratante": "nombre", "contratista": "nombre"}},
  "objeto": "objeto aqui",
  "resumen": "resumen aqui",
  "riesgo_general": "alto",
  "firmar": false,
  "puntos_criticos": ["punto1", "punto2"],
  "recomendaciones": ["rec1", "rec2"]
}}
</formato_salida>
"""
    try:
        response = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=2048,
            temperature=0,
            system=CONTRACTAI_SYSTEM,
            messages=[{"role": "user", "content": prompt}]
        )
        return limpiar_json(response.content[0].text)
    except Exception as e:
        print(f"  ❌ Error al consolidar: {e}")
        return None


# ── Analizar contrato completo ──
def analizar_contrato(ruta_pdf):
    print(f"\n📄 Leyendo PDF...")
    paginas = leer_pdf(ruta_pdf)
    print(f"  ✅ {len(paginas)} páginas leídas")

    # Documento corto — analizar directo

    texto_completo = " ".join([p["texto"] for p in paginas])
    texto_completo = texto_completo.replace('"', "'")

    if len(texto_completo) <= 50000:
        print("  Analizando documento completo...")
        try:
            response = client.messages.create(
                model="claude-sonnet-4-5",
                max_tokens=4096,
                temperature=0,
                system=CONTRACTAI_SYSTEM,
                messages=[{"role": "user", "content": 
                          prompt_analisis_contrato(texto_completo)}]
            )
            ok, error = verificar_respuesta(response)
            if not ok:
                print(f"  ❌ {error}")
                return None
            return limpiar_json(response.content[0].text)
        except Exception as e:
            print(f"  ❌ Error: {e}")
            return None

    # Documento largo — dividir en chunks
    print(f"  Documento largo — dividiendo en partes...")
    chunks = dividir_en_chunks(paginas)
    print(f"  ✅ {len(chunks)} partes generadas")

    resultados = []
    for i, chunk in enumerate(chunks):
        resultado = analizar_chunk(chunk, i + 1, len(chunks))
        if not resultado["exitoso"]:
            print(f"\n  ❌ ERROR en parte {resultado['parte']}: {resultado['error']}")
            print("  ⛔ Análisis detenido.")
            return None
        print(f"  ✅ Parte {i+1} analizada")
        resultados.append(resultado)

    return consolidar(resultados)

# ── Generar reporte Word ──
def generar_reporte(datos, ruta_salida):
    doc = Document()

    titulo = doc.add_heading("ContractAI — Reporte de Análisis", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fecha = doc.add_paragraph(
        f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading("Datos del documento", 1)
    tabla = doc.add_table(rows=4, cols=2)
    tabla.style = "Table Grid"
    datos_tabla = [
        ["Tipo",        datos.get("tipo_documento", "—")],
        ["Contratante", datos.get("partes", {}).get("contratante", "—")],
        ["Contratista", datos.get("partes", {}).get("contratista", "—")],
        ["Objeto",      datos.get("objeto", "—")],
    ]
    for i, (campo, valor) in enumerate(datos_tabla):
        tabla.rows[i].cells[0].text = campo
        tabla.rows[i].cells[1].text = valor
        tabla.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_heading("Resumen ejecutivo", 1)
    doc.add_paragraph(datos.get("resumen", "—"))
    doc.add_paragraph()

    doc.add_heading("Decisión", 1)
    riesgo = datos.get("riesgo_general", "—").upper()
    firma = "✅ SE PUEDE FIRMAR" if datos.get("firmar") else "❌ NO FIRMAR TAL CUAL"
    p = doc.add_paragraph()
    run = p.add_run(f"Riesgo: {riesgo}   |   {firma}")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_heading("Puntos críticos", 1)
    for punto in datos.get("puntos_criticos", []):
        doc.add_paragraph(punto, style="List Bullet")

    doc.add_paragraph()
    doc.add_heading("Recomendaciones", 1)
    for rec in datos.get("recomendaciones", []):
        doc.add_paragraph(rec, style="List Bullet")

    doc.add_paragraph()
    pie = doc.add_paragraph(
        "Generado por ContractAI — Herramienta de análisis O&G"
    )
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.runs[0].font.size = Pt(8)
    pie.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.save(ruta_salida)
    print(f"  ✅ Reporte guardado: {ruta_salida}")

# ── Mostrar resultado en consola ──
def mostrar_resultado(datos):
    separador("ANÁLISIS — ContractAI")
    print(f"Tipo      : {datos.get('tipo_documento', '—')}")
    print(f"Riesgo    : {datos.get('riesgo_general', '—').upper()}")
    firma = "✅ Se puede firmar" if datos.get("firmar") else "❌ NO firmar tal cual"
    print(f"Decisión  : {firma}")
    print(f"\nResumen   : {datos.get('resumen', '—')}")

    print("\n⚠️  Puntos críticos:")
    for p in datos.get("puntos_criticos", []):
        print(f"  • {p}")

    print("\n✅ Recomendaciones:")
    for r in datos.get("recomendaciones", []):
        print(f"  • {r}")