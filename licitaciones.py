import anthropic
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz

client = anthropic.Anthropic()

SYSTEM_LICITACIONES = """
Sos un consultor senior especializado en licitaciones de Oil & Gas 
en Argentina con 15 años de experiencia en Vaca Muerta, Neuquén.

Conocés en profundidad:
- Pliegos de licitación de YPF, Shell, TotalEnergies y Pampa Energía
- Requisitos técnicos y financieros estándar del sector
- Prácticas comerciales de empresas proveedoras de la cuenca neuquina
- Marco regulatorio: Ley 17.319, normativas IAPG y Secretaría de Energía
- Factores que determinan el éxito en procesos licitatorios petroleros

TUS REGLAS:
- Respondés ÚNICAMENTE en JSON sin texto adicional
- Nunca inventás datos — si falta información lo indicás claramente
- Tus estimaciones de probabilidad son realistas, no optimistas
- Siempre considerás el contexto del mercado neuquino
- Usás SOLO comillas simples dentro de los valores de texto del JSON
"""

def leer_pdf(ruta):
    """Lee el contenido de un PDF."""
    try:
        doc = fitz.open(ruta)
        texto = ""
        for pagina in doc:
            texto += pagina.get_text()
        doc.close()
        texto = texto.replace('"', "'").replace('\\', ' ').replace('\x00', '')
        return texto[:60000]
    except Exception as e:
        return f"Error al leer PDF: {str(e)}"

def analizar_licitacion(texto, nombre_archivo):
    """Analiza una licitación individual."""

    prompt = f"""
<licitacion>
Archivo: {nombre_archivo}
{texto}
</licitacion>

<instruccion>
Analizá esta licitación de Oil & Gas y extraé toda la información 
relevante para evaluar si conviene presentar oferta.
</instruccion>

<formato_salida>
{{
  "nombre_archivo": "{nombre_archivo}",
  "titulo": "",
  "organismo_licitante": "",
  "tipo_licitacion": "",
  "objeto": "",
  "monto_estimado_usd": 0,
  "moneda": "",
  "plazo_ejecucion_dias": 0,
  "fecha_apertura": "",
  "requisitos_tecnicos": [],
  "requisitos_financieros": {{
    "patrimonio_neto_min_usd": 0,
    "facturacion_min_usd": 0,
    "garantia_oferta_usd": 0,
    "otros": []
  }},
  "clausulas_riesgo": [
    {{"clausula": "", "nivel": "alto/medio/bajo"}}
  ],
  "puntos_favor": [],
  "puntos_contra": [],
  "competencia_estimada": "alta/media/baja",
  "complejidad_tecnica": "alta/media/baja",
  "score_oportunidad": 0,
  "recomendacion": "presentar/evaluar/no_presentar",
  "resumen_ejecutivo": ""
}}
</formato_salida>
"""

    response = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=2048,
        temperature=0,
        system=SYSTEM_LICITACIONES,
        messages=[{"role": "user", "content": prompt}]
    )

    texto_resp = response.content[0].text
    texto_resp = texto_resp.replace("```json", "").replace("```", "").strip()
    match = re.search(r'\{.*\}', texto_resp, re.DOTALL)
    if not match:
        return None
    try:
        return json.loads(match.group())
    except json.JSONDecodeError:
        try:
            import re as re2
            json_limpio = re2.sub(
                r':\s*"([^"]*)"',
                lambda m: ': "' + m.group(1).replace('\n', ' ').replace('\r', '') + '"',
                match.group()
            )
            return json.loads(json_limpio)
        except:
            return None

def comparar_licitaciones(analisis_lista):
    """Compara todas las licitaciones y genera el ranking final."""

    resumen = json.dumps(analisis_lista, ensure_ascii=False, indent=2)

    prompt = f"""
<analisis_licitaciones>
{resumen}
</analisis_licitaciones>

<instruccion>
Comparás todas las licitaciones analizadas y generás un ranking 
definitivo indicando cuál conviene más presentar para una PyME 
proveedora de servicios en Vaca Muerta con capacidad media.

El score_oportunidad va de 0 a 100 donde:
- 80-100: Muy recomendada — alta probabilidad de ganar
- 60-79:  Recomendada — buena oportunidad
- 40-59:  Evaluar — depende de la capacidad de la empresa
- 0-39:   No recomendada — riesgo alto o baja probabilidad

Considerá:
- Relación monto/complejidad
- Requisitos accesibles para una PyME
- Nivel de competencia esperado
- Riesgo contractual
- Plazo de ejecución realista

Usás SOLO comillas simples dentro de los valores de texto.
</instruccion>

<formato_salida>
{{
  "fecha_analisis": "{datetime.now().strftime('%d/%m/%Y %H:%M')}",
  "total_licitaciones": {len(analisis_lista)},
  "ranking": [
    {{
      "posicion": 1,
      "nombre_archivo": "",
      "titulo": "",
      "score": 0,
      "recomendacion": "presentar/evaluar/no_presentar",
      "monto_estimado_usd": 0,
      "probabilidad_ganar": "alta/media/baja",
      "motivo_ranking": "",
      "ventajas_clave": [],
      "riesgos_clave": []
    }}
  ],
  "recomendacion_general": "",
  "mejor_opcion": "",
  "advertencias": []
}}
</formato_salida>
"""

    response = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=2048,
        temperature=0,
        system=SYSTEM_LICITACIONES,
        messages=[{"role": "user", "content": prompt}]
    )

    texto_resp = response.content[0].text
    texto_resp = texto_resp.replace("```json", "").replace("```", "").strip()
    match = re.search(r'\{.*\}', texto_resp, re.DOTALL)
    if not match:
        return None
    try:
        return json.loads(match.group())
    except json.JSONDecodeError:
        try:
            import re as re2
            json_limpio = re2.sub(
                r':\s*"([^"]*)"',
                lambda m: ': "' + m.group(1).replace('\n', ' ').replace('\r', '') + '"',
                match.group()
            )
            return json.loads(json_limpio)
        except:
            return None

def generar_word_comparativo(analisis_lista, comparacion, ruta_salida):
    """Genera el reporte Word comparativo."""

    doc = Document()

    titulo = doc.add_heading("ANÁLISIS COMPARATIVO DE LICITACIONES", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run(
        f"Oil & Gas — Vaca Muerta, Neuquén   |   "
        f"{comparacion.get('fecha_analisis', '—')}"
    )
    run.font.size = Pt(11)
    run.font.bold = True

    doc.add_paragraph()

    # Resumen ejecutivo
    doc.add_heading("Resumen Ejecutivo", 1)
    total = comparacion.get("total_licitaciones", 0)
    mejor = comparacion.get("mejor_opcion", "—")
    doc.add_paragraph(
        f"Se analizaron {total} licitación(es). "
        f"La opción más recomendada es: {mejor}."
    )
    doc.add_paragraph(comparacion.get("recomendacion_general", "—"))

    # Advertencias
    advertencias = comparacion.get("advertencias", [])
    if advertencias:
        doc.add_paragraph()
        doc.add_heading("⚠️ Advertencias", 2)
        for adv in advertencias:
            doc.add_paragraph(adv, style="List Bullet")

    doc.add_paragraph()

    # Ranking
    doc.add_heading("Ranking de Licitaciones", 1)

    medallas = {1: "🥇", 2: "🥈", 3: "🥉"}
    colores_rec = {
        "presentar": RGBColor(40, 167, 69),
        "evaluar": RGBColor(255, 193, 7),
        "no_presentar": RGBColor(220, 53, 69)
    }

    for item in comparacion.get("ranking", []):
        pos = item.get("posicion", 0)
        medalla = medallas.get(pos, f"#{pos}")

        doc.add_heading(
            f"{medalla} Posición {pos}: {item.get('titulo', item.get('nombre_archivo', '—'))}",
            2
        )

        # Tabla de datos clave
        tbl = doc.add_table(rows=5, cols=2)
        tbl.style = "Table Grid"
        campos = [
            ["Score de oportunidad", f"{item.get('score', 0)}/100"],
            ["Recomendación", item.get("recomendacion", "—").upper()],
            ["Monto estimado", f"USD {item.get('monto_estimado_usd', 0):,.0f}"],
            ["Probabilidad de ganar", item.get("probabilidad_ganar", "—").upper()],
            ["Archivo", item.get("nombre_archivo", "—")],
        ]
        for i, (campo, valor) in enumerate(campos):
            tbl.rows[i].cells[0].text = campo
            tbl.rows[i].cells[1].text = valor
            tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            # Color según recomendación en la fila de recomendación
            if campo == "Recomendación":
                rec = item.get("recomendacion", "evaluar")
                color = colores_rec.get(rec, RGBColor(0, 0, 0))
                tbl.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = color
                tbl.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()
        doc.add_paragraph(f"Motivo: {item.get('motivo_ranking', '—')}")

        ventajas = item.get("ventajas_clave", [])
        if ventajas:
            doc.add_paragraph("Ventajas clave:").runs[0].bold = True
            for v in ventajas:
                doc.add_paragraph(v, style="List Bullet")

        riesgos = item.get("riesgos_clave", [])
        if riesgos:
            doc.add_paragraph("Riesgos clave:").runs[0].bold = True
            for r in riesgos:
                doc.add_paragraph(r, style="List Bullet")

        doc.add_paragraph()

    # Detalle por licitación
    doc.add_heading("Detalle por Licitación", 1)

    for analisis in analisis_lista:
        doc.add_heading(
            analisis.get("titulo", analisis.get("nombre_archivo", "—")),
            2
        )
        doc.add_paragraph(analisis.get("resumen_ejecutivo", "—"))

        req_tec = analisis.get("requisitos_tecnicos", [])
        if req_tec:
            doc.add_paragraph("Requisitos técnicos:").runs[0].bold = True
            for r in req_tec:
                doc.add_paragraph(r, style="List Bullet")

        doc.add_paragraph()

    # Pie
    pie = doc.add_paragraph(
        "Reporte generado por OilAI — Módulo de Licitaciones · "
        "Vaca Muerta, Neuquén · 2026"
    )
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.runs[0].font.size = Pt(8)
    pie.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.save(ruta_salida)