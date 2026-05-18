import streamlit as st
import anthropic
import json
import re
from datetime import date
import io

# ── Word ────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── PDF ─────────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT


# ═══════════════════════════════════════════════════════════════════════════
# 1. GENERACIÓN DE ÍTEMS CON IA
# ═══════════════════════════════════════════════════════════════════════════

def generar_items_presupuesto(descripcion, empresa_emisora, empresa_cliente, lugar):
    client = anthropic.Anthropic()

    prompt = f"""Sos un experto en presupuestos técnico-comerciales para empresas de servicios 
en la industria Oil & Gas de Argentina (Cuenca Neuquina / Vaca Muerta).

Empresa que emite: {empresa_emisora}
Empresa cliente:   {empresa_cliente}
Lugar/Yacimiento:  {lugar}
Descripción del trabajo: {descripcion}

Generá un presupuesto detallado dividido en DOS secciones:
1. MANO DE OBRA – categoría del operario, tarea, cantidad, unidad (jornal/hora/día)
2. MATERIALES   – descripción del material/insumo, cantidad, unidad (metro/kg/unidad/etc.)

Respondé ÚNICAMENTE con JSON válido, sin texto adicional ni backticks.
Estructura exacta:

{{
  "numero_presupuesto": "P-2025-001",
  "descripcion_trabajo": "descripción resumida y profesional del trabajo",
  "mano_de_obra": [
    {{
      "item": 1,
      "descripcion": "descripción del ítem de mano de obra",
      "cantidad": 10,
      "unidad": "jornal",
      "precio_unitario": 0
    }}
  ],
  "materiales": [
    {{
      "item": 1,
      "descripcion": "descripción del material",
      "cantidad": 100,
      "unidad": "metro",
      "precio_unitario": 0
    }}
  ],
  "notas": "Condiciones de pago sugeridas, validez, observaciones técnicas relevantes."
}}"""

    try:
        message = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=2000,
            temperature=0,
            messages=[{"role": "user", "content": prompt}]
        )
        texto = message.content[0].text
        match = re.search(r'\{.*\}', texto, re.DOTALL)
        if match:
            return json.loads(match.group())
        else:
            st.error(f"⚠️ La IA no devolvió JSON válido. Respuesta: {texto[:300]}")
    except Exception as e:
        st.error(f"⚠️ Error real: {str(e)}")
    return None


# ═══════════════════════════════════════════════════════════════════════════
# 2. CÁLCULOS
# ═══════════════════════════════════════════════════════════════════════════

def calcular_totales(data):
    total_mo  = sum(i['cantidad'] * i['precio_unitario'] for i in data['mano_de_obra'])
    total_mat = sum(i['cantidad'] * i['precio_unitario'] for i in data['materiales'])
    return total_mo, total_mat, total_mo + total_mat


# ═══════════════════════════════════════════════════════════════════════════
# 3. EXPORTAR A WORD
# ═══════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color):
    """Aplica color de fondo a una celda de tabla en Word."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _fila_header(tabla, encabezados, color_bg='1F4E79'):
    fila = tabla.rows[0]
    for idx, txt in enumerate(encabezados):
        c = fila.cells[idx]
        c.text = ''
        run = c.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(c, color_bg)


def generar_word(data, empresa_emisora, empresa_cliente, lugar, fecha, validez, moneda):
    doc = Document()

    # Márgenes
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── Título ──────────────────────────────────────────────────────────────
    titulo = doc.add_heading('PRESUPUESTO', level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"N° {data['numero_presupuesto']}   |   Fecha: {fecha}   |   Validez: {validez} días")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ── Cabecera: emisor / cliente / lugar / moneda ─────────────────────────
    t = doc.add_table(rows=2, cols=2)
    t.style = 'Table Grid'

    celdas = [
        (0, 0, 'PROVEEDOR',    empresa_emisora),
        (0, 1, 'CLIENTE',      empresa_cliente),
        (1, 0, 'LUGAR / YACIMIENTO', lugar),
        (1, 1, 'MONEDA',       moneda),
    ]
    for fila, col, label, valor in celdas:
        c = t.cell(fila, col)
        c.text = ''
        c.paragraphs[0].add_run(f"{label}\n").bold = True
        c.paragraphs[0].add_run(valor)
        c.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── Descripción del trabajo ──────────────────────────────────────────────
    doc.add_heading('Descripción del Trabajo', level=2)
    p = doc.add_paragraph(data['descripcion_trabajo'])
    p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    total_mo, total_mat, total_general = calcular_totales(data)

    headers = ['Ítem', 'Descripción', 'Cant.', 'Unidad', f'P. Unit. ({moneda})', f'Total ({moneda})']

    # ── Tabla Mano de Obra ──────────────────────────────────────────────────
    doc.add_heading('MANO DE OBRA', level=2)

    tabla_mo = doc.add_table(rows=1 + len(data['mano_de_obra']) + 1, cols=6)
    tabla_mo.style = 'Table Grid'
    _fila_header(tabla_mo, headers)

    for n, item in enumerate(data['mano_de_obra'], start=1):
        fila  = tabla_mo.rows[n]
        total = item['cantidad'] * item['precio_unitario']
        for col_idx, val in enumerate([
            str(item['item']), item['descripcion'], str(item['cantidad']),
            item['unidad'], f"{item['precio_unitario']:,.2f}", f"{total:,.2f}"
        ]):
            fila.cells[col_idx].text = val
            fila.cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9)

    # Subtotal MO
    fila_st = tabla_mo.rows[-1]
    _set_cell_bg(fila_st.cells[4], 'D9E2F3')
    _set_cell_bg(fila_st.cells[5], 'D9E2F3')
    c4 = fila_st.cells[4]
    c4.text = ''
    c4.paragraphs[0].add_run('SUBTOTAL M.O.').bold = True
    c4.paragraphs[0].runs[0].font.size = Pt(9)
    c5 = fila_st.cells[5]
    c5.text = ''
    c5.paragraphs[0].add_run(f"{total_mo:,.2f}").bold = True
    c5.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── Tabla Materiales ────────────────────────────────────────────────────
    doc.add_heading('MATERIALES', level=2)

    tabla_mat = doc.add_table(rows=1 + len(data['materiales']) + 1, cols=6)
    tabla_mat.style = 'Table Grid'
    _fila_header(tabla_mat, headers)

    for n, item in enumerate(data['materiales'], start=1):
        fila  = tabla_mat.rows[n]
        total = item['cantidad'] * item['precio_unitario']
        for col_idx, val in enumerate([
            str(item['item']), item['descripcion'], str(item['cantidad']),
            item['unidad'], f"{item['precio_unitario']:,.2f}", f"{total:,.2f}"
        ]):
            fila.cells[col_idx].text = val
            fila.cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9)

    fila_st = tabla_mat.rows[-1]
    _set_cell_bg(fila_st.cells[4], 'D9E2F3')
    _set_cell_bg(fila_st.cells[5], 'D9E2F3')
    c4 = fila_st.cells[4]
    c4.text = ''
    c4.paragraphs[0].add_run('SUBTOTAL MAT.').bold = True
    c4.paragraphs[0].runs[0].font.size = Pt(9)
    c5 = fila_st.cells[5]
    c5.text = ''
    c5.paragraphs[0].add_run(f"{total_mat:,.2f}").bold = True
    c5.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── Total General ────────────────────────────────────────────────────────
    t_total = doc.add_table(rows=1, cols=2)
    t_total.style = 'Table Grid'
    _set_cell_bg(t_total.cell(0, 0), '1F4E79')
    _set_cell_bg(t_total.cell(0, 1), '1F4E79')
    for idx, txt in enumerate([f'TOTAL GENERAL ({moneda})', f'{total_general:,.2f}']):
        c = t_total.cell(0, idx)
        c.text = ''
        run = c.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    # ── Notas ────────────────────────────────────────────────────────────────
    if data.get('notas'):
        doc.add_heading('Notas y Condiciones', level=2)
        doc.add_paragraph(data['notas']).runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── Firma ────────────────────────────────────────────────────────────────
    p_firma = doc.add_paragraph('_' * 45)
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre = doc.add_paragraph(empresa_emisora)
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre.runs[0].bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════════════════════
# 4. EXPORTAR A PDF
# ═══════════════════════════════════════════════════════════════════════════

def generar_pdf(data, empresa_emisora, empresa_cliente, lugar, fecha, validez, moneda):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2*cm, leftMargin=2.5*cm,
        topMargin=2*cm, bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()
    AZUL   = colors.HexColor('#1F4E79')
    AZUL_L = colors.HexColor('#D9E2F3')
    BLANCO = colors.white
    GRIS   = colors.HexColor('#555555')

    estilo_titulo = ParagraphStyle('titulo', fontSize=20, textColor=AZUL,
                                   alignment=TA_CENTER, spaceAfter=4, fontName='Helvetica-Bold')
    estilo_sub    = ParagraphStyle('sub', fontSize=9, textColor=GRIS,
                                   alignment=TA_CENTER, spaceAfter=12)
    estilo_h2     = ParagraphStyle('h2', fontSize=11, textColor=AZUL,
                                   fontName='Helvetica-Bold', spaceBefore=12, spaceAfter=4)
    estilo_cuerpo = ParagraphStyle('cuerpo', fontSize=9, spaceAfter=6)
    estilo_nota   = ParagraphStyle('nota', fontSize=8, textColor=GRIS, spaceAfter=6)

    story = []
    total_mo, total_mat, total_general = calcular_totales(data)

    # Título
    story.append(Paragraph('PRESUPUESTO', estilo_titulo))
    story.append(Paragraph(
        f"N° {data['numero_presupuesto']}   |   Fecha: {fecha}   |   Validez: {validez} días",
        estilo_sub
    ))
    story.append(HRFlowable(width='100%', thickness=1, color=AZUL))
    story.append(Spacer(1, 0.3*cm))

    # Cabecera
    cabecera = [
        [Paragraph('<b>PROVEEDOR</b>', estilo_nota),   Paragraph('<b>CLIENTE</b>', estilo_nota)],
        [Paragraph(empresa_emisora, estilo_cuerpo),     Paragraph(empresa_cliente, estilo_cuerpo)],
        [Paragraph('<b>LUGAR / YACIMIENTO</b>', estilo_nota), Paragraph('<b>MONEDA</b>', estilo_nota)],
        [Paragraph(lugar, estilo_cuerpo),               Paragraph(moneda, estilo_cuerpo)],
    ]
    t_cab = Table(cabecera, colWidths=[8*cm, 8*cm])
    t_cab.setStyle(TableStyle([
        ('BOX',        (0,0), (-1,-1), 0.5, AZUL),
        ('INNERGRID',  (0,0), (-1,-1), 0.3, colors.HexColor('#CCCCCC')),
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F5F8FC')),
        ('VALIGN',     (0,0), (-1,-1), 'MIDDLE'),
        ('LEFTPADDING',(0,0), (-1,-1), 6),
        ('BOTTOMPADDING',(0,0),(-1,-1),4),
    ]))
    story.append(t_cab)
    story.append(Spacer(1, 0.4*cm))

    # Descripción
    story.append(Paragraph('Descripción del Trabajo', estilo_h2))
    story.append(Paragraph(data['descripcion_trabajo'], estilo_cuerpo))

    # Función auxiliar para tablas de ítems
    def tabla_items(items_list):
        header = [
            Paragraph('<b>Ítem</b>', estilo_nota),
            Paragraph('<b>Descripción</b>', estilo_nota),
            Paragraph('<b>Cant.</b>', estilo_nota),
            Paragraph('<b>Unidad</b>', estilo_nota),
            Paragraph(f'<b>P. Unit. ({moneda})</b>', estilo_nota),
            Paragraph(f'<b>Total ({moneda})</b>', estilo_nota),
        ]
        filas = [header]
        subtotal = 0
        for item in items_list:
            total_item = item['cantidad'] * item['precio_unitario']
            subtotal  += total_item
            filas.append([
                Paragraph(str(item['item']),    estilo_cuerpo),
                Paragraph(item['descripcion'],  estilo_cuerpo),
                Paragraph(str(item['cantidad']),estilo_cuerpo),
                Paragraph(item['unidad'],       estilo_cuerpo),
                Paragraph(f"{item['precio_unitario']:,.2f}", estilo_cuerpo),
                Paragraph(f"{total_item:,.2f}", estilo_cuerpo),
            ])
        # Fila subtotal
        filas.append([
            '', '', '', '',
            Paragraph('<b>SUBTOTAL</b>', estilo_nota),
            Paragraph(f'<b>{subtotal:,.2f}</b>', estilo_nota),
        ])
        t = Table(filas, colWidths=[1*cm, 6.5*cm, 1.5*cm, 1.5*cm, 2.5*cm, 2.5*cm])
        t.setStyle(TableStyle([
            ('BACKGROUND',   (0,0),  (-1,0),  AZUL),
            ('TEXTCOLOR',    (0,0),  (-1,0),  BLANCO),
            ('BACKGROUND',   (0,-1), (-1,-1), AZUL_L),
            ('BOX',          (0,0),  (-1,-1), 0.5, AZUL),
            ('INNERGRID',    (0,0),  (-1,-1), 0.3, colors.HexColor('#CCCCCC')),
            ('VALIGN',       (0,0),  (-1,-1), 'MIDDLE'),
            ('LEFTPADDING',  (0,0),  (-1,-1), 4),
            ('BOTTOMPADDING',(0,0),  (-1,-1), 4),
            ('ROWBACKGROUNDS',(0,1), (-1,-2), [colors.white, colors.HexColor('#F5F8FC')]),
        ]))
        return t

    # Mano de Obra
    story.append(Paragraph('MANO DE OBRA', estilo_h2))
    story.append(tabla_items(data['mano_de_obra']))
    story.append(Spacer(1, 0.4*cm))

    # Materiales
    story.append(Paragraph('MATERIALES', estilo_h2))
    story.append(tabla_items(data['materiales']))
    story.append(Spacer(1, 0.5*cm))

    # Total General
    t_gen = Table(
        [[Paragraph(f'<b>TOTAL GENERAL ({moneda})</b>', ParagraphStyle('tg', fontSize=12, textColor=BLANCO, fontName='Helvetica-Bold')),
          Paragraph(f'<b>{total_general:,.2f}</b>',      ParagraphStyle('tv', fontSize=12, textColor=BLANCO, fontName='Helvetica-Bold', alignment=TA_RIGHT))]],
        colWidths=[11*cm, 5*cm]
    )
    t_gen.setStyle(TableStyle([
        ('BACKGROUND',   (0,0), (-1,-1), AZUL),
        ('BOX',          (0,0), (-1,-1), 0.5, AZUL),
        ('LEFTPADDING',  (0,0), (-1,-1), 8),
        ('RIGHTPADDING', (0,0), (-1,-1), 8),
        ('BOTTOMPADDING',(0,0), (-1,-1), 8),
        ('TOPPADDING',   (0,0), (-1,-1), 8),
    ]))
    story.append(t_gen)
    story.append(Spacer(1, 0.4*cm))

    # Notas
    if data.get('notas'):
        story.append(Paragraph('Notas y Condiciones', estilo_h2))
        story.append(Paragraph(data['notas'], estilo_nota))
        story.append(Spacer(1, 0.5*cm))

    # Firma
    story.append(HRFlowable(width='45%', thickness=0.7, color=AZUL, hAlign='CENTER'))
    story.append(Paragraph('Servicios OIL IA', ParagraphStyle('firma', fontSize=9,
                            alignment=TA_CENTER, fontName='Helvetica-Bold')))
    story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════════════════════
# 5. PÁGINA PRINCIPAL DE STREAMLIT
# ═══════════════════════════════════════════════════════════════════════════

def presupuestos_page():
    st.title("💰 PresupuestosAI")
    st.markdown("Generá presupuestos técnico-comerciales profesionales para trabajos en Oil & Gas.")

    st.info(
        "💡 **Versión demo:** ingresás los precios manualmente. "
        "En la **versión oficial**, los campos de precio unitario serán editables "
        "directamente en pantalla antes de exportar.",
        icon="ℹ️"
    )

    # ── Formulario de entrada ────────────────────────────────────────────────
    with st.form("form_presupuesto"):
        st.subheader("📋 Datos del Presupuesto")
        col1, col2 = st.columns(2)

        with col1:
            empresa_emisora = st.text_input("Empresa que emite el presupuesto *",
                                            placeholder="Ej: Buen Aike Servicios S.R.L.")
            empresa_cliente = st.text_input("Empresa cliente",
                                            placeholder="Ej: YPF S.A.")
            lugar           = st.text_input("Lugar / Yacimiento",
                                            placeholder="Ej: Loma Campana, Neuquén")

        with col2:
            fecha   = st.date_input("Fecha", value=date.today())
            validez = st.number_input("Validez del presupuesto (días)",
                                      min_value=1, max_value=90, value=15)
            moneda  = st.selectbox("Moneda", ["USD", "ARS"])

        st.subheader("📝 Descripción del Trabajo")
        descripcion = st.text_area(
            "Describí el trabajo a presupuestar *",
            height=140,
            placeholder=(
                "Ej: Soldadura de piping 6 pulgadas acero A106 Grado B, 200 metros lineales. "
                "Incluye prefabricado en taller, arenado Sa 2½ y pintura epoxi poliuretano 2 manos. "
                "Trabajo en yacimiento Loma Campana, Neuquén."
            )
        )

        submitted = st.form_submit_button("🤖 Generar Presupuesto con IA",
                                          use_container_width=True, type="primary")

    if submitted:
        if not descripcion.strip() or not empresa_emisora.strip():
            st.warning("⚠️ Completá al menos la empresa emisora y la descripción del trabajo.")
            return

        with st.spinner("🔄 Analizando el trabajo y generando ítems con IA..."):
            data = generar_items_presupuesto(descripcion, empresa_emisora, empresa_cliente, lugar)

        if not data:
            st.error("❌ No se pudo generar el presupuesto. Revisá tu API key o intentá de nuevo.")
            return

        st.session_state['presupuesto_data']   = data
        st.session_state['presupuesto_params'] = {
            'empresa_emisora': empresa_emisora,
            'empresa_cliente': empresa_cliente,
            'lugar':   lugar,
            'fecha':   str(fecha),
            'validez': validez,
            'moneda':  moneda,
        }
        st.success("✅ Presupuesto generado. Ingresá los precios unitarios y exportá.")

    # ── Visualización + carga de precios ─────────────────────────────────────
    if 'presupuesto_data' not in st.session_state:
        return

    data   = st.session_state['presupuesto_data']
    params = st.session_state['presupuesto_params']
    moneda = params['moneda']

    st.markdown("---")
    st.subheader(f"📄 Presupuesto N° {data['numero_presupuesto']}")
    st.markdown(f"**{data['descripcion_trabajo']}**")
    st.caption(f"Cliente: {params['empresa_cliente']}  |  Lugar: {params['lugar']}  |  Validez: {params['validez']} días")

    # ── Mano de Obra ─────────────────────────────────────────────────────────
    st.markdown("### 👷 Mano de Obra")
    st.caption("⚙️ *Versión oficial: precios editables en tabla. Demo: ingresalos aquí.*")

    col_hdrs = st.columns([0.5, 3.5, 1, 1, 2, 1.8])
    for h, lbl in zip(col_hdrs, ['Ítem', 'Descripción', 'Cant.', 'Unidad',
                                   f'P. Unit. ({moneda})', f'Total ({moneda})']):
        h.markdown(f"**{lbl}**")

    for i, item in enumerate(data['mano_de_obra']):
        cols = st.columns([0.5, 3.5, 1, 1, 2, 1.8])
        cols[0].write(item['item'])
        cols[1].write(item['descripcion'])
        cols[2].write(item['cantidad'])
        cols[3].write(item['unidad'])
        precio = cols[4].number_input(
            label=f"mo_{i}", label_visibility="collapsed",
            min_value=0.0, value=float(item['precio_unitario']),
            step=500.0, key=f"mo_{i}"
        )
        data['mano_de_obra'][i]['precio_unitario'] = precio
        cols[5].write(f"{precio * item['cantidad']:,.2f}")

    total_mo = sum(i['cantidad'] * i['precio_unitario'] for i in data['mano_de_obra'])
    st.markdown(f"**Subtotal Mano de Obra: {moneda} {total_mo:,.2f}**")

    # ── Materiales ────────────────────────────────────────────────────────────
    st.markdown("### 🔩 Materiales")
    st.caption("⚙️ *Versión oficial: precios editables en tabla. Demo: ingresalos aquí.*")

    col_hdrs = st.columns([0.5, 3.5, 1, 1, 2, 1.8])
    for h, lbl in zip(col_hdrs, ['Ítem', 'Descripción', 'Cant.', 'Unidad',
                                   f'P. Unit. ({moneda})', f'Total ({moneda})']):
        h.markdown(f"**{lbl}**")

    for i, item in enumerate(data['materiales']):
        cols = st.columns([0.5, 3.5, 1, 1, 2, 1.8])
        cols[0].write(item['item'])
        cols[1].write(item['descripcion'])
        cols[2].write(item['cantidad'])
        cols[3].write(item['unidad'])
        precio = cols[4].number_input(
            label=f"mat_{i}", label_visibility="collapsed",
            min_value=0.0, value=float(item['precio_unitario']),
            step=500.0, key=f"mat_{i}"
        )
        data['materiales'][i]['precio_unitario'] = precio
        cols[5].write(f"{precio * item['cantidad']:,.2f}")

    total_mat = sum(i['cantidad'] * i['precio_unitario'] for i in data['materiales'])
    st.markdown(f"**Subtotal Materiales: {moneda} {total_mat:,.2f}**")

    # ── Total General ─────────────────────────────────────────────────────────
    total_general = total_mo + total_mat
    st.markdown("---")
    st.markdown(
        f"<h2 style='color:#1F4E79;'>💰 TOTAL GENERAL: {moneda} {total_general:,.2f}</h2>",
        unsafe_allow_html=True
    )

    if data.get('notas'):
        st.info(f"📌 {data['notas']}")

    # ── Botones de exportación ────────────────────────────────────────────────
    st.markdown("### 📥 Exportar")
    col1, col2 = st.columns(2)

    with col1:
        word_buf = generar_word(
            data,
            params['empresa_emisora'], params['empresa_cliente'],
            params['lugar'],           params['fecha'],
            params['validez'],         moneda
        )
        st.download_button(
            label="📄 Descargar Word (.docx)",
            data=word_buf,
            file_name=f"Presupuesto_{data['numero_presupuesto']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with col2:
        pdf_buf = generar_pdf(
            data,
            params['empresa_emisora'], params['empresa_cliente'],
            params['lugar'],           params['fecha'],
            params['validez'],         moneda
        )
        st.download_button(
            label="📑 Descargar PDF",
            data=pdf_buf,
            file_name=f"Presupuesto_{data['numero_presupuesto']}.pdf",
            mime="application/pdf",
            use_container_width=True
        )
